#encoding=utf-8
from docx import Document

foreword = 'Foreword'
document_read = Document()

#0是完整目录，1是一级目录
titleList = []
paraList = []
Property = []

#用于批量加载
def run(fileName):
    global document_read
    document_read = Document(fileName)
    articleProperty_TSTR()
    Content0()
    Para()
    # testOut()
    List = Out()
    return List

#读取文章属性
def articleProperty_TSTR():
    global Property
    Property = []
    text = document_read.paragraphs[0]
    text = text.text.strip()
    text = text.split(' ')
    str = 'standardType' + '$' + text[0]
    Property.append(str)
    str = 'docType' + '$' + text[1]
    Property.append(str)
    str = '_id' + '$' + text[2]+ text[3]
    Property.append(str)
    str = 'type' + '$' + text[2]
    Property.append(str)
    str = 'Version' + '$' + text[3]
    Property.append(str)
    #print(Property)
    return

#读取整个目录
def Content0():
    judge = 0
    #0代表没有遇见目录，1代表遇见目录，2代表已经遍历完目录
    global titleList
    titleList = []
    for text in document_read.paragraphs:
       para = text.text
       para = para.strip()
       if judge == 0:
           if foreword in para:
               #加入Forword
              pos = para.rfind('	')
              if pos != -1:
                  para = para[0:pos]
              para = para.strip()
              titleList.append(para)
              judge = 1
              continue
       if judge == 1:
            if foreword in para:
                return
            else:
                pos = para.rfind('	')
                if pos != -1:
                    para = para[0: pos]
                para = para.strip()
                titleList.append(para)
                continue
       # if judge == 2:
       #   paraList.append(para)
       #   continue
    return

#比较目录文字与正文的文字
def compare(s1,s2):
    s1 = s1.replace(' ','')
    s1 = s1.replace('\t','')
    s1 = s1.replace('\n','')
    s2 = s2.replace(' ', '')
    s2 = s2.replace('\t', '')
    s2 = s2.replace('\n','')
    if s1 == s2:
        return 1
    return 0

#测试直接输出文本
def testOut():
    for i in range(0,len(titleList) -1 ):
        print("TITLE:" + titleList[i])
        for l in paraList[i]:
            print(l)
    return

#输出分解好的文本
def Out():
    List = []
    for i in Property:
        List.append(i)
        List.append('\n')
    for i in range(0,len(titleList) - 1 ):
        combineList = []
        text = []
        combineList.append(titleList[i])
        for l in paraList[i]:
            text.append(l)
        str0 = ''.join(text)
        combineList.append(str0)
        str = "$".join(combineList)
        List.append(str)
        List.append('\n')
    return List

#拆分段落
def Para():
    judge = 0
    line = 0
    key = 1
    global paraList
    paraList = []
    paraPart = []
    length = len(titleList)
    #0代表没有遇见目录，1代表遇见目录，2代表已经遍历完目录
    for text in document_read.paragraphs:
       para = text.text
       para = para.strip()
       if para.isspace():
           continue
      # print(para)
       if judge == 0:
           if foreword in para:
              judge = 1
           continue
       elif judge == 1:
            if foreword in para:
                judge = 2
       elif judge == 2:
           #print(titleList[key] + '   is   ' + para)
           if compare(para,titleList[key]):
               paraList.append(paraPart)
               paraPart = []
               if key < length - 1:
                  key = key + 1
           else:
               paraPart.append(para)
    paraList.append(paraPart)
    return

#用于批量输出到word
def outWord(List, fileName):
    # document_write = Document()
    # document_write.add_paragraph([paragraph  for paragraph in List ])
    # document_write.save(fileName)
    f1 = open(fileName, 'w',encoding='utf-8')
    for paragraph in List:
        f1.write(paragraph)
    f1.close()


