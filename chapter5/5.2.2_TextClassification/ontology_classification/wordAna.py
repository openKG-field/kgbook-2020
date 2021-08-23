# -*- encoding: utf-8 -*-
"""
@File    : wordAna.py
@Time    : 2021/3/19 14:00
@Author  : Ryt_tech
@Email   : zhaohongyu2401@yeah.net
@Software: PyCharm
"""

import docx
from docx import Document


def _fill_blank(table):
    cols = max([len(i) for i in table])
 
    new_table = []
    for i, row in enumerate(table):
        new_row = []
        [new_row.extend([i] * int(cols / len(row))) for i in row]
        print(new_row)
        new_table.append(new_row)
 
    return new_table

import re


e = 'economic.docx'
count = 1
document = Document(e)

for t in  document.tables:
    for i in range(len(t.rows)):
        for j in range(len(t.columns)):
            print (str(j) + ' : ',t.cell(i,j).text,)
        print ('')

# for p in document.paragraphs:
#     text = p.text
#     if text.isspace() or len(text) < 1 :
#         continue
#     if re.findall(r'[a-zA-Z]', text) and len(text) < 25:
#         print text
#         count += 1
# print count 
    
    